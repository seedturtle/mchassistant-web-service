"""
MCH Assistant Web Service - 門諾醫院AI語音助理
語音轉文字並生成專業報告（支援分段錄音、Word模板）
"""

import os
import io
import re
import uuid
import logging
import base64
import tempfile
import json
import urllib.parse
from datetime import datetime
from typing import Optional, List
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor, TimeoutError

from fastapi import FastAPI, HTTPException, Request, UploadFile, File, Form
from fastapi.responses import HTMLResponse, RedirectResponse, Response
from fastapi.staticfiles import StaticFiles
from fastapi.responses import JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from docx import Document
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText

# =============================================================================
# Configuration
# =============================================================================

ACCESS_PASSWORD = "ABC1234"
UPLOAD_DIR = Path("./uploads")
UPLOAD_DIR.mkdir(exist_ok=True)
RECORDING_DIR = Path("./recordings")
RECORDING_DIR.mkdir(exist_ok=True)

MATON_API_KEY = os.getenv("MATON_API_KEY", "")
DRIVE_UPLOAD_URL = "https://gateway.maton.ai/google-drive/upload/drive/v3/files"
DRIVE_FOLDER_ID = "15goCYQxn8xM7R1HoLLTS-Hbv-lSge8Y2"  # 吞嚥障礙問卷篩檢報告
DRIVE_TEMPLATES_FOLDER_ID = "1VCeYlNLRwVfp7rZnrbwwKDXNZwZcd95K"  # 報告類型模板
DRIVE_AUTO_UPLOAD_ID = "1HRRcjWcL4r4CCw958GTgcqunysi5QhwU"  # 報告自動上傳資料夾
GMAIL_GATEWAY = "https://gateway.maton.ai/google-mail/gmail/v1/users/me/messages/send"
REPORT_TYPES_FILE = Path("./report_types.json")

HF_TOKEN = os.getenv("HF_TOKEN", "")
MINIMAX_API_KEY = os.getenv("MINIMAX_API_KEY", "")
MINIMAX_API_GATEWAY = "https://api.minimax.io/v1/text/chatcompletion_v2"

sessions = {}
executor = ThreadPoolExecutor(max_workers=2)

# =============================================================================
# Dynamic Report Types (persisted to JSON)
# =============================================================================

DEFAULT_REPORT_TYPES = {
    "general": "一般報告",
    "medical": "醫療報告",
    "meeting": "會議記錄",
    "swallow": "吞嚥評估",
    "ent": "耳鼻喉科報告"
}

def load_report_types() -> dict:
    if REPORT_TYPES_FILE.exists():
        try:
            return json.loads(REPORT_TYPES_FILE.read_text())
        except Exception as e:
            logging.warning(f"Failed to load {REPORT_TYPES_FILE}: {e}")
    data = dict(DEFAULT_REPORT_TYPES)
    save_report_types(data)
    return data

def save_report_types(data: dict):
    REPORT_TYPES_FILE.write_text(json.dumps(data, ensure_ascii=False, indent=2))

report_types_store = load_report_types()

TEMPLATES_DIR = Path("./templates")
TEMPLATES_DIR.mkdir(exist_ok=True)

# =============================================================================
# Pydantic Models
# =============================================================================

class AudioSegmentRequest(BaseModel):
    audio_data: str  # Base64 encoded
    format: str = "webm"

class EmailRequest(BaseModel):
    to_email: str
    subject: str = "MCH Assistant 報告"
    body: str = ""

class DysphagiaUploadRequest(BaseModel):
    report: str
    patientName: str
    date: str
    to: str = ""  # optional, for email endpoint

# =============================================================================
# Lifespan
# =============================================================================

async def lifespan(app: FastAPI):
    logging.info("MCH Assistant Web Service 啟動中...")
    logging.info("正在同步 Google Drive 模板...")
    _sync_templates_from_drive()
    yield
    logging.info("MCH Assistant Web Service 關閉中...")

# =============================================================================
# FastAPI Application
# =============================================================================

app = FastAPI(
    title="MCH Assistant 語音助理",
    description="門諾醫院AI語音助理 - 語音轉文字生成報告",
    version="2.2.0",
    lifespan=lifespan
)

# CORS: allow dysphagia screening frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "https://mchdysphagiascreen.zeabur.app",
        "http://localhost:8080",
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

BASE_DIR = Path(__file__).parent
app.mount("/static", StaticFiles(directory=BASE_DIR / "static"), name="static")

# =============================================================================
# Helper Functions
# =============================================================================

def get_session_id(request: Request) -> Optional[str]:
    return request.cookies.get("session_id")

def validate_session(request: Request) -> bool:
    session_id = get_session_id(request)
    return session_id and session_id in sessions

def get_or_create_session() -> str:
    session_id = str(uuid.uuid4())[:8]
    sessions[session_id] = {
        "report_type": "general",
        "segments": [],            # final transcribed text segments
        "audio_files": {},         # {uid: {filename, filepath, ext, source: "upload"|"record"}}
        "generated_docx": None,
        "created_at": datetime.now(),
        "auto_email": "",
        "auto_email_sent": False,
        "auto_email_message": None,
        "auto_email_error": None,
        "drive_folder_id": "",
        "drive_upload_result": None,
        "processing": False,
        "processing_done": False,
        "processing_error": None,
        "processing_progress": None,  # {transcribed, total, current_file, stage}
    }
    return session_id

def transcribe_audio(audio_bytes: bytes, file_ext: str = ".webm") -> str:
    try:
        from faster_whisper import WhisperModel
        import os as os_module
        if HF_TOKEN:
            os_module.environ["HF_TOKEN"] = HF_TOKEN
        model = WhisperModel("small", device="cpu", compute_type="int8")
        with tempfile.NamedTemporaryFile(suffix=file_ext, delete=False) as f:
            f.write(audio_bytes)
            temp_path = f.name
        try:
            segments, info = model.transcribe(temp_path, language="zh")
            text = "".join([s.text for s in segments]).strip()
            return text if text else "[無辨識結果]"
        finally:
            os.unlink(temp_path)
    except Exception as e:
        logging.error(f"Faster Whisper transcription error: {e}")
        return f"[轉換失敗: {str(e)}]"

def summarize_with_hermes(transcribed_text: str, report_type: str, placeholders: list = None) -> any:
    logging.info(f"[summarize_with_hermes] report_type={report_type}, text_length={len(transcribed_text)}")
    if not MINIMAX_API_KEY:
        if placeholders:
            return {p: transcribed_text for p in placeholders}
        return transcribed_text
    
    try:
        import urllib.request
        import urllib.error
        
        type_name = report_types_store.get(report_type, report_type)
        
        if placeholders:
            fields_str = "、".join(placeholders)
            user_prompt = f"""報告類型：{type_name}

模板欄位：{fields_str}

錄音內容：
{transcribed_text}

請根據錄音內容，為以上每個模板欄位產生合適的文字內容。
以 JSON 格式回傳，不要有其他文字。範例：
{{"{placeholders[0]}": "填寫內容", "{placeholders[-1]}": "填寫內容"}}"""
            system_prompt = "你是專業的醫療報告整理助理。你只回傳純 JSON，不附加任何說明文字。"
        else:
            system_prompt = "你是專業的醫療報告整理助理。請將下面的口語錄音整理成正式格式，直接輸出結果，不需要標記【段落】。"
            user_prompt = f"報告類型：{type_name}\n\n錄音內容：\n{transcribed_text}"
        
        payload = json.dumps({
            "model": "minimax-m2.7",
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            "temperature": 0.3,
            "max_tokens": 4000
        }).encode()
        
        req = urllib.request.Request(MINIMAX_API_GATEWAY, data=payload, method='POST')
        req.add_header('Authorization', f'Bearer {MINIMAX_API_KEY}')
        req.add_header('Content-Type', 'application/json')
        
        with urllib.request.urlopen(req, timeout=60) as resp:
            result = json.loads(resp.read().decode())
            if 'choices' in result and len(result['choices']) > 0:
                content = result['choices'][0]['message']['content'].strip()
                if placeholders:
                    try:
                        clean = content.strip()
                        if clean.startswith("```"):
                            clean = clean.split("\n", 1)[1] if "\n" in clean else clean
                            clean = clean.rsplit("```", 1)[0] if "```" in clean else clean
                            clean = clean.strip()
                        fields_result = json.loads(clean)
                        for p in placeholders:
                            if p not in fields_result:
                                fields_result[p] = f"[待補充：{p}]"
                        return fields_result
                    except (json.JSONDecodeError, Exception) as e:
                        return {p: content if i == 0 else f"[待補充：{p}]" for i, p in enumerate(placeholders)}
                else:
                    return content
            if placeholders:
                return {p: f"[待補充：{p}]" for p in placeholders}
            return transcribed_text
    except Exception as e:
        logging.error(f"MiniMax summarization error: {e}")
        if placeholders:
            return {p: transcribed_text for p in placeholders}
        return transcribed_text

def fill_template(template_path: str, segments: list, report_type: str, summarized_text: str = None, fields_dict: dict = None) -> bytes:
    doc = Document(template_path)
    if summarized_text:
        full_text = summarized_text
    else:
        full_text = "\n\n".join([f"【段落{i+1}】\n{seg['transcription']}" for i, seg in enumerate(segments)])
    report_date = datetime.now().strftime('%Y年%m月%d日')
    report_name = get_report_type_name(report_type)
    
    def replace_paragraph_text(para):
        full_para_text = "".join(run.text for run in para.runs)
        new_text = full_para_text
        nonlocal content_replaced
        if "{{content}}" in new_text:
            new_text = new_text.replace("{{content}}", full_text)
            content_replaced = True
        if "{{date}}" in new_text:
            new_text = new_text.replace("{{date}}", report_date)
        if "{{report_type}}" in new_text:
            new_text = new_text.replace("{{report_type}}", report_name)
        if fields_dict:
            for key, value in fields_dict.items():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in new_text:
                    new_text = new_text.replace(placeholder, str(value))
        if new_text != full_para_text:
            for i, run in enumerate(para.runs):
                if i == 0:
                    run.text = new_text
                else:
                    run.text = ""
    
    content_replaced = False
    for para in doc.paragraphs:
        replace_paragraph_text(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_paragraph_text(para)
    if not content_replaced and full_text:
        doc.add_paragraph("")
        doc.add_paragraph(full_text)
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

def get_report_type_name(report_type: str) -> str:
    return report_types_store.get(report_type, report_type)

def extract_placeholders(doc_path: str) -> list:
    doc = Document(doc_path)
    placeholders = set()
    for para in doc.paragraphs:
        placeholders.update(re.findall(r'\{\{(.*?)\}\}', para.text))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                placeholders.update(re.findall(r'\{\{(.*?)\}\}', cell.text))
    auto_fill = {"date", "report_type", "content"}
    return [p for p in placeholders if p not in auto_fill]

def _send_email_sync(session: dict, to_email: str) -> tuple:
    try:
        import urllib.request, urllib.error
        report_date = datetime.now().strftime('%Y年%m月%d日')
        report_type_name = get_report_type_name(session.get("report_type", "general"))
        msg = MIMEMultipart('mixed')
        msg['to'] = to_email
        msg['subject'] = f"MCH {report_type_name} - {report_date}"
        body = f"""您好，

這是由 MCH Assistant 產生的 {report_type_name}，日期：{report_date}。

彙整報告已作為 Word 檔案 (.docx) 附加於此郵件中，請直接下載開啟。

---
此郵件由 MCH Assistant 自動發送
"""
        msg.attach(MIMEText(body, 'plain', 'utf-8'))
        if session.get("generated_docx"):
            from email.mime.base import MIMEBase
            from email import encoders
            clean_date = report_date.replace("年", "").replace("月", "").replace("日", "")
            part = MIMEBase('application', 'vnd.openxmlformats-officedocument.wordprocessingml.document')
            part.set_payload(session["generated_docx"])
            encoders.encode_base64(part)
            part.add_header('Content-Disposition', f'attachment; filename="MCH_{report_type_name}_{clean_date}.docx"')
            msg.attach(part)
        raw = base64.urlsafe_b64encode(msg.as_bytes()).decode().rstrip('=')
        data = json.dumps({"raw": raw}).encode()
        gmail_req = urllib.request.Request(GMAIL_GATEWAY, data=data, method='POST')
        gmail_req.add_header('Authorization', f'Bearer {MATON_API_KEY}')
        gmail_req.add_header('Content-Type', 'application/json')
        with urllib.request.urlopen(gmail_req, timeout=30) as resp:
            return True, f"報告已自動寄送至 {to_email}"
    except Exception as e:
        return False, str(e)

# =============================================================================
# Background Processing Pipeline
# =============================================================================

def _process_generate(session_id: str):
    """Background pipeline: transcribe all → summarize → Word → Email"""
    session = sessions.get(session_id)
    if not session:
        return
    
    session["processing"] = True
    session["processing_done"] = False
    session["segments"] = []
    
    audio_files = session.get("audio_files", {})
    total = len(audio_files)
    session["processing_progress"] = {"stage": "transcribing", "transcribed": 0, "total": total, "current_file": "", "error": None}
    
    try:
        # --- Step 1: Transcribe all audio files ---
        for uid, info in sorted(audio_files.items(), key=lambda x: x[1].get("created_at", 0)):
            session["processing_progress"]["current_file"] = info["filename"]
            try:
                with open(info["filepath"], "rb") as f:
                    audio_bytes = f.read()
                text = transcribe_audio(audio_bytes, file_ext=info["ext"])
            except Exception as e:
                text = f"[轉換失敗: {str(e)}]"
            
            session["segments"].append({
                "id": uid,
                "transcription": text,
                "audio_path": info["filename"]
            })
            session["processing_progress"]["transcribed"] += 1
        
        # --- Step 2: Summarize with MiniMax ---
        session["processing_progress"]["stage"] = "summarizing"
        full_text = "\n\n".join([f"【段落{i+1}】\n{seg['transcription']}" for i, seg in enumerate(session["segments"])])
        report_type = session.get("report_type", "general")
        
        # Use type-specific template, fallback to generic
        type_template = TEMPLATES_DIR / f"{report_type}.docx"
        generic_template = TEMPLATES_DIR / "template.docx"
        if type_template.exists():
            template_path = str(type_template)
        elif generic_template.exists():
            template_path = str(generic_template)
        else:
            template_path = None
        
        template_fields = []
        template_usable = False
        if template_path and Path(template_path).exists():
            try:
                template_fields = extract_placeholders(template_path)
                template_usable = True
            except Exception:
                pass
        
        fields_dict = None
        summarized_text = None
        
        if MINIMAX_API_KEY:
            result = summarize_with_hermes(full_text, report_type, placeholders=template_fields if template_fields else None)
            if template_fields:
                fields_dict = result
            else:
                summarized_text = result
        else:
            summarized_text = full_text
        
        # --- Step 3: Generate Word ---
        session["processing_progress"]["stage"] = "generating"
        if template_usable:
            docx_bytes = fill_template(template_path, session["segments"], report_type,
                                       summarized_text=summarized_text, fields_dict=fields_dict)
            session["generated_docx"] = docx_bytes
        
        # --- Step 4: Send email if configured ---
        target_email = session.get("auto_email", "")
        if target_email:
            session["processing_progress"]["stage"] = "emailing"
            success, msg = _send_email_sync(session, target_email)
            if success:
                session["auto_email_sent"] = True
                session["auto_email_message"] = msg
            else:
                session["auto_email_error"] = msg
        
        # --- Step 5: Auto upload report to Google Drive ---
        if session.get("generated_docx") and MATON_API_KEY:
            # User can override the auto folder; otherwise use default
            drive_folder = session.get("drive_folder_id", "") or DRIVE_AUTO_UPLOAD_ID
            session["processing_progress"]["stage"] = "drive_upload"
            try:
                now = datetime.now()
                filename = now.strftime('%Y%m%d_%H%M%S') + ".docx"
                success, result = _upload_binary_to_drive(
                    session["generated_docx"],
                    filename,
                    drive_folder
                )
                session["drive_upload_result"] = {
                    "success": success,
                    "file_id": result if success else None,
                    "error": None if success else result,
                    "filename": filename
                }
            except Exception as e:
                session["drive_upload_result"] = {
                    "success": False, "file_id": None, "error": str(e), "filename": ""
                }
        
        session["processing_progress"]["stage"] = "done"
        session["processing_done"] = True
        
    except Exception as e:
        logging.error(f"Background generate error: {e}")
        session["processing_error"] = str(e)
        if session.get("processing_progress"):
            session["processing_progress"]["error"] = str(e)
            session["processing_progress"]["stage"] = "error"
        session["processing_done"] = True
    finally:
        session["processing"] = False


# =============================================================================
# API Endpoints
# =============================================================================

# --- Upload audio files (save only, no transcription) ---
@app.post("/api/sessions/{session_id}/upload")
async def upload_audio_files(
    session_id: str,
    request: Request,
    files: List[UploadFile] = File(...),
    email: str = Form("")
):
    if not validate_session(request) or get_session_id(request) != session_id:
        raise HTTPException(status_code=401, detail="Not authenticated")
    if session_id not in sessions:
        raise HTTPException(status_code=404, detail="Session not found")
    if not files:
        return {"success": False, "error": "請選擇至少一個音檔"}
    
    if email:
        sessions[session_id]["auto_email"] = email
        sessions[session_id]["auto_email_sent"] = False
        sessions[session_id]["auto_email_message"] = None
        sessions[session_id]["auto_email_error"] = None
    
    allowed_extensions = {".wav", ".mp3", ".m4a", ".ogg", ".webm", ".flac", ".aac", ".wma", ".opus"}
    max_bytes = 100 * 1024 * 1024
    results = []
    
    for file in files:
        ext = Path(file.filename).suffix.lower() if file.filename else ".webm"
        if ext not in allowed_extensions:
            results.append({"filename": file.filename, "success": False, "error": f"不支援的格式: {ext}"})
            continue
        content = await file.read()
        if len(content) > max_bytes:
            results.append({"filename": file.filename, "success": False, "error": f"檔案過大: {len(content)/(1024*1024):.1f}MB，上限 100MB"})
            continue
        
        uid = uuid.uuid4().hex[:8]
        safe_name = f"{session_id}_{uid}{ext}"
        filepath = UPLOAD_DIR / safe_name
        with open(filepath, "wb") as f:
            f.write(content)
        
        sessions[session_id]["audio_files"][uid] = {
            "filename": file.filename, "filepath": str(filepath), "ext": ext,
            "size": len(content), "source": "upload", "created_at": datetime.now().timestamp()
        }
        results.append({"uid": uid, "filename": file.filename, "success": True})
    
    return {"success": True, "results": results, "auto_email": email if email else None}


# --- Record audio segment (save only, no transcription) ---
@app.post("/api/sessions/{session_id}/segment")
async def add_segment(session_id: str, request: Request, seg: AudioSegmentRequest):
    if not validate_session(request) or get_session_id(request) != session_id:
        raise HTTPException(status_code=401, detail="Not authenticated")
    if session_id not in sessions:
        raise HTTPException(status_code=404, detail="Session not found")
    
    audio_bytes = base64.b64decode(seg.audio_data)
    uid = uuid.uuid4().hex[:8]
    ext = f".{seg.format}" if seg.format else ".webm"
    safe_name = f"{session_id}_record_{uid}{ext}"
    filepath = RECORDING_DIR / safe_name
    with open(filepath, "wb") as f:
        f.write(audio_bytes)
    
    sessions[session_id]["audio_files"][uid] = {
        "filename": f"錄音段落{list(sessions[session_id]['audio_files'].keys()).count('') + 1}",
        "filepath": str(filepath), "ext": ext,
        "size": len(audio_bytes), "source": "record", "created_at": datetime.now().timestamp()
    }
    # Give it a friendly name
    idx = len([v for v in sessions[session_id]["audio_files"].values() if v["source"] == "record"])
    sessions[session_id]["audio_files"][uid]["filename"] = f"錄音段落 {idx}"
    
    return {"success": True, "segment_id": uid}


# --- Get list of audio files in session ---
@app.get("/api/sessions/{session_id}/audio-files")
async def get_audio_files(session_id: str, request: Request):
    if not validate_session(request) or get_session_id(request) != session_id:
        raise HTTPException(status_code=401, detail="Not authenticated")
    if session_id not in sessions:
        raise HTTPException(status_code=404, detail="Session not found")
    return {
        "success": True,
        "files": sessions[session_id].get("audio_files", {}),
        "total": len(sessions[session_id].get("audio_files", {}))
    }


# --- Generate report (start background pipeline) ---
@app.post("/api/sessions/{session_id}/generate")
async def generate_report(session_id: str, request: Request, req: dict):
    if not validate_session(request) or get_session_id(request) != session_id:
        raise HTTPException(status_code=401, detail="Not authenticated")
    if session_id not in sessions:
        raise HTTPException(status_code=404, detail="Session not found")
    
    session = sessions[session_id]
    audio_files = session.get("audio_files", {})
    if not audio_files:
        return {"success": False, "error": "尚無音檔，請先錄音或上傳"}
    
    if session.get("processing", False):
        return {"success": False, "error": "已有處理程序正在執行"}
    
    report_type = req.get("report_type", "general")
    session["report_type"] = report_type
    
    # Reset processing state
    session["processing"] = True
    session["processing_done"] = False
    session["processing_error"] = None
    session["processing_progress"] = {"stage": "queued", "transcribed": 0, "total": len(audio_files), "current_file": ""}
    
    # Submit to background
    executor.submit(_process_generate, session_id)
    
    return {"success": True, "message": f"背景處理已啟動，共 {len(audio_files)} 個音檔", "total": len(audio_files)}


# --- Get generation progress ---
@app.get("/api/sessions/{session_id}/generate-status")
async def get_generate_status(session_id: str, request: Request):
    if not validate_session(request) or get_session_id(request) != session_id:
        raise HTTPException(status_code=401, detail="Not authenticated")
    if session_id not in sessions:
        raise HTTPException(status_code=404, detail="Session not found")
    
    s = sessions[session_id]
    return {
        "success": True,
        "processing": s.get("processing", False),
        "processing_done": s.get("processing_done", False),
        "processing_error": s.get("processing_error"),
        "processing_progress": s.get("processing_progress"),
        "segments_count": len(s.get("segments", [])),
        "has_docx": s.get("generated_docx") is not None,
        "auto_email": s.get("auto_email") or None,
        "auto_email_sent": s.get("auto_email_sent", False),
        "auto_email_message": s.get("auto_email_message"),
        "auto_email_error": s.get("auto_email_error"),
        "drive_folder_id": s.get("drive_folder_id") or None,
        "drive_upload_result": s.get("drive_upload_result")
    }


# =============================================================================
# Pages (HTML)
# =============================================================================

@app.get("/", response_class=HTMLResponse)
async def root(request: Request):
    if not validate_session(request):
        return RedirectResponse(url="/login", status_code=302)
    return RedirectResponse(url="/dashboard", status_code=302)

@app.get("/login", response_class=HTMLResponse)
async def login_page(request: Request):
    if validate_session(request):
        return RedirectResponse(url="/dashboard", status_code=302)
    html = """<!DOCTYPE html>
<html lang="zh-TW">
<head>
    <meta charset="UTF-8"><meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>MCH Assistant - 登入</title><link rel="stylesheet" href="/static/css/style.css">
</head>
<body>
    <div class="login-container">
        <h1>🏥 MCH Assistant</h1><p class="subtitle">門諾醫院 AI 語音助理</p>
        <p id="error" class="error" style="display:none">密碼錯誤，請重新輸入</p>
        <form id="loginForm">
            <div class="input-group"><label for="password">🔐 請輸入存取密碼</label>
            <input type="password" id="password" name="password" placeholder="輸入密碼" required></div>
            <button type="submit" class="btn">登入</button>
        </form>
        <p class="info">僅限醫院內部人員使用</p>
    </div>
    <script>
    document.getElementById('loginForm').onsubmit = async (e) => {
        e.preventDefault();
        const pwd = document.getElementById('password').value;
        const res = await fetch('/api/auth/login', {method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify({password:pwd})});
        if (res.ok) window.location.href='/dashboard';
        else document.getElementById('error').style.display='block';
    };
    </script>
</body></html>"""
    return HTMLResponse(content=html)

@app.get("/dashboard", response_class=HTMLResponse)
async def dashboard(request: Request):
    if not validate_session(request):
        return RedirectResponse(url="/login", status_code=302)
    session_id = get_session_id(request)
    
    html = f"""<!DOCTYPE html>
<html lang="zh-TW">
<head>
    <meta charset="UTF-8"><meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>MCH Assistant - 儀表板</title><link rel="stylesheet" href="/static/css/style.css">
</head>
<body>
    <div class="header">
        <div class="logo">🏥 MCH Assistant</div>
        <div><span class="session-id">Session: {session_id}</span><a href="/new" class="logout">🔄 新開工作</a></div>
    </div>
    <div class="container">
        <h1>🎙️ 語音報告助理</h1>
        <p class="subtitle">錄音或上傳音檔，完成後按下「🔄 批次產出報告」</p>
        
        <div class="card instructions-card">
            <h3>📖 使用說明</h3>
            <ul class="instructions-list">
                <li>🎤 按錄音鈕開始，再按一下停止，段落會自動列出</li>
                <li>📁 也可上傳音檔（MP3/WAV/M4A 等），支援拖放或多選，每檔上限 100MB</li>
                <li>🔄 錄音與上傳的音檔會累積在列表中，可混合使用</li>
                <li>⚡ 按下「即時處理」在頁面上觀看進度，完成後手動下載或寄送</li>
                <li>📧 按下「背景處理並寄送」需先輸入 Email，完成後自動寄送到信箱，可關閉網頁</li>
                <li>📥 處理完成後也可手動下載或按「手動寄送」</li>
                <li>📄 可上傳 Word 模板（{{content}}、{{date}}、{{report_type}}），AI 內容自動填入</li>
            </ul>
        </div>
        
        <div class="card">
            <h3>📋 報告類型</h3>
            <div style="display:flex; gap:8px; align-items:stretch;">
                <select id="reportType" class="select-full" style="flex:1;">
                    <option value="" disabled selected>載入中...</option>
                </select>
                <button id="manageTypesBtn" class="btn btn-secondary" style="white-space:nowrap; padding:12px 16px; font-size:18px; line-height:1; border-radius:10px;" title="管理報告類型">✏️</button>
            </div>
        </div>
        
        <div class="card">
            <h3>📄 Word 模板（選填）</h3>
            <p class="hint">使用 {{content}}、{{date}}、{{report_type}} 作為佔位符</p>
            <p class="hint" style="font-size:12px; color:#4ade80;">📎 模板會依所選報告類型儲存，並同步上傳 Google Drive</p>
            <div id="templateStatus" class="template-status"></div>
            <div id="templateTypeStatus" class="template-status" style="margin-top:8px;"></div>
            <input type="file" id="templateFile" accept=".docx" class="file-input">
        </div>
        
        <div class="card recorder-card">
            <h3>🎤 即時錄音</h3>
            <div class="recorder-box">
                <button class="btn-record" id="recordBtn">🎤</button>
                <div class="status" id="status">點擊麥克風開始錄音</div>
            </div>
        </div>
        
        <div class="card upload-card">
            <h3>📁 上傳音檔</h3>
            <p class="hint">支援格式：MP3、WAV、M4A、OGG、FLAC、AAC、WebM、Opus ｜ 每檔上限 100MB</p>
            <div class="upload-box" id="uploadBox">
                <div class="upload-icon">📂</div>
                <div class="upload-text">點擊選擇檔案 或 拖放音檔到此處</div>
                <div class="upload-hint">可選擇多個檔案一次上傳</div>
                <input type="file" id="fileInput" multiple accept=".wav,.mp3,.m4a,.ogg,.webm,.flac,.aac,.wma,.opus" class="file-input-hidden">
            </div>
        </div>
        
        <div class="card">
            <h3>📋 音檔列表 <span id="fileCount" class="file-count">(0)</span></h3>
            <div id="audioFileList" class="audio-file-list"></div>
            <div class="clear-btn-wrap"><button class="btn btn-secondary" id="clearBtn">🗑 清空重置</button></div>
        </div>
        
        <div id="processingArea" class="card processing-card" style="display:none">
            <h3>⚙️ 批次處理進度</h3>
            <div class="progress-bar"><div class="progress-fill" id="genProgressFill"></div></div>
            <div class="progress-text" id="genProgressText">準備中...</div>
            <div id="genSegments" class="segments-container" style="max-height:200px;margin-top:12px"></div>
            <div id="genResult" style="margin-top:8px"></div>
        </div>
        
        <div class="actions">
            <div class="email-input-group">
                <label for="emailInput">📧 背景模式 Email（必填）：按「背景處理並寄送」時，完成後自動寄至此信箱</label>
                <input type="email" id="emailInput" placeholder="example@mch.org.tw" class="input-full">
            </div>
            <div class="email-input-group">
                <label for="driveFolderInput">☁️ Google Drive 資料夾 ID（選填）：預設自動上傳到指定資料夾，可在此更改上傳目標</label>
                <div style="display:flex; gap:8px;">
                    <input type="text" id="driveFolderInput" placeholder="留空則用預設資料夾" class="input-full" style="flex:1;">
                    <button id="setDriveFolderBtn" class="btn btn-secondary" style="padding:12px 20px; white-space:nowrap; border-radius:10px;">設定</button>
                </div>
                <div id="driveFolderStatus" class="template-status"></div>
            </div>
            <button class="btn btn-email-mode" id="bgEmailBtn">📧 背景處理並寄送</button>
            <button class="btn btn-generate" id="generateBtn">⚡ 即時處理</button>
            <button class="btn btn-download" id="downloadBtn" disabled>📥 下載報告</button>
            <button class="btn btn-email" id="emailBtn" disabled>📧 手動寄送</button>
        </div>
        <div id="driveResult" class="result-box" style="display:none;"></div>
        
        <div id="result" class="result-box"></div>
    </div>
    
    <!-- Report Type Management Modal -->
    <div id="typeModal" class="modal-overlay" style="display:none;">
        <div class="modal-content">
            <div class="modal-header">
                <h3>📋 管理報告類型</h3>
                <button id="closeModalBtn" class="modal-close">✕</button>
            </div>
            <div class="modal-body">
                <div class="type-list" id="typeList"></div>
                <div class="type-add">
                    <input type="text" id="newTypeName" class="input-full" placeholder="輸入新報告類型名稱..." maxlength="30">
                    <button id="addTypeBtn" class="btn btn-secondary" style="margin-top:8px; width:100%;">＋ 新增</button>
                </div>
            </div>
        </div>
    </div>
    
    <script>
    let mediaRecorder, audioChunks = [], isRecording = false;
    let SESSION_ID = '{session_id}';
    let recordingTimer = null;
    let genPollTimer = null;
    let isProcessing = false;

    function escapeHtml(t) {{ const d = document.createElement('div'); d.textContent = t; return d.innerHTML; }}

    // Refresh audio file list from server
    async function refreshFileList() {{
        try {{
            const res = await fetch('/api/sessions/' + SESSION_ID + '/audio-files');
            const data = await res.json();
            if (!data.success) return;
            const files = data.files || {{}};
            const countEl = document.getElementById('fileCount');
            if (countEl) countEl.textContent = '(' + Object.keys(files).length + ')';
            const listEl = document.getElementById('audioFileList');
            if (!listEl) return;
            let html = '';
            for (const [uid, info] of Object.entries(files)) {{
                const icon = info.source === 'record' ? '🎤' : '📁';
                const sizeKB = (info.size / 1024).toFixed(1);
                html += '<div class="audio-file-item">' +
                    '<span class="audio-file-icon">' + icon + '</span>' +
                    '<span class="audio-file-name">' + escapeHtml(info.filename) + '</span>' +
                    '<span class="audio-file-size">' + sizeKB + 'KB</span>' +
                    '<button class="btn-delete-file" onclick="deleteAudioFile(\\'' + uid + '\\')" title="刪除此音檔">✕</button>' +
                    '</div>';
            }}
            listEl.innerHTML = html || '<div class="audio-file-empty">尚無音檔，請錄音或上傳</div>';
        }} catch(e) {{}}
    }}

    function disableInputsDuringProcessing(disabled) {{
        document.getElementById('recordBtn').disabled = disabled;
        document.getElementById('uploadBox').style.pointerEvents = disabled ? 'none' : 'auto';
        document.getElementById('uploadBox').style.opacity = disabled ? '0.4' : '1';
        document.getElementById('fileInput').disabled = disabled;
        document.getElementById('emailInput').disabled = disabled;
        document.getElementById('templateFile').disabled = disabled;
        document.getElementById('bgEmailBtn').disabled = disabled;
        document.getElementById('generateBtn').disabled = disabled;
    }}

    // Poll generation status
    async function pollGenerateStatus() {{
        try {{
            const res = await fetch('/api/sessions/' + SESSION_ID + '/generate-status');
            const data = await res.json();
            if (!data.success) return;
            
            const progress = data.processing_progress || {{}};
            const total = progress.total || 0;
            const transcribed = progress.transcribed || 0;
            const stage = progress.stage || '';
            
            let pct = 0;
            let stageLabel = '';
            if (stage === 'transcribing') {{
                pct = total > 0 ? Math.round(transcribed / total * 70) : 0;
                stageLabel = '🔄 語音辨識中 ' + transcribed + '/' + total;
            }} else if (stage === 'summarizing') {{
                pct = 75;
                stageLabel = '🤖 AI 彙整中';
            }} else if (stage === 'generating') {{
                pct = 90;
                stageLabel = '📄 產出 Word 報告中';
            }} else if (stage === 'emailing') {{
                pct = 90;
                stageLabel = '📧 寄送 Email 中';
            }} else if (stage === 'drive_upload') {{
                pct = 95;
                stageLabel = '☁️ 上傳 Google Drive 中';
            }} else if (stage === 'done') {{
                pct = 100;
                stageLabel = '✅ 完成';
            }} else if (stage === 'error') {{
                pct = 0;
                stageLabel = '❌ 處理失敗';
            }}
            
            document.getElementById('genProgressFill').style.width = pct + '%';
            document.getElementById('genProgressText').textContent = stageLabel;
            
            if (data.processing_done) {{
                if (genPollTimer) {{ clearInterval(genPollTimer); genPollTimer = null; }}
                isProcessing = false;
                
                if (data.processing_error) {{
                    document.getElementById('genResult').innerHTML = '<div class="error">❌ ' + escapeHtml(data.processing_error) + '</div>';
                }} else {{
                    if (data.auto_email_sent) {{
                        document.getElementById('genResult').innerHTML = '<div class="success">✅ 報告已自動寄送至 ' + escapeHtml(data.auto_email) + '</div>';
                    }} else if (data.auto_email) {{
                        document.getElementById('genResult').innerHTML = '<div class="warning">⚠️ 已產生報告，但自動寄信失敗：' + escapeHtml(data.auto_email_error || '') + '</div>';
                    }} else {{
                        document.getElementById('genResult').innerHTML = '<div class="success">✅ 報告已產生！</div>';
                    }}
                    // Show Drive upload result
                    if (data.drive_upload_result) {{
                        const dr = data.drive_upload_result;
                        const driveDiv = document.getElementById('driveResult');
                        if (dr.success) {{
                            driveDiv.style.display = 'block';
                            driveDiv.innerHTML = '<div class="success">☁️ 報告已上傳 Google Drive：' + escapeHtml(dr.filename) + '</div>';
                        }} else if (dr.error) {{
                            driveDiv.style.display = 'block';
                            driveDiv.innerHTML = '<div class="warning">⚠️ Drive 上傳失敗：' + escapeHtml(dr.error) + '</div>';
                        }}
                    }}
                    document.getElementById('downloadBtn').disabled = false;
                    document.getElementById('emailBtn').disabled = false;
                }}
                disableInputsDuringProcessing(false);
            }}
        }} catch(e) {{}}
    }}

    // ========== Template Upload ==========
    document.getElementById('templateFile').addEventListener('change', async () => {{
        const file = document.getElementById('templateFile').files[0];
        if (!file) return;
        const rt = document.getElementById('reportType');
        const fd = new FormData();
        fd.append('file', file);
        fd.append('session_id', SESSION_ID);
        fd.append('report_type', rt ? rt.value : 'template');
        const res = await fetch('/api/sessions/' + SESSION_ID + '/template', {{method:'POST', body:fd}});
        const data = await res.json();
        const typeName = rt && rt.options[rt.selectedIndex] ? rt.options[rt.selectedIndex].textContent : '';
        document.getElementById('templateStatus').innerHTML = data.success 
            ? '<span class="success">✓ 已上傳：' + data.filename + '（' + typeName + '）</span>' 
            : '<span class="error">上傳失敗</span>';
        if (data.success && data.drive_status) {{
            document.getElementById('templateStatus').innerHTML += '<br><span style="font-size:12px; color:#4ade80;">' + data.drive_status + '</span>';
        }}
    }});

    // ========== Recording ==========
    document.getElementById('recordBtn').addEventListener('click', async function() {{
        if (isProcessing) {{ alert('處理中，無法錄音'); return; }}
        if (!isRecording) {{
            try {{
                const stream = await navigator.mediaDevices.getUserMedia({{audio:true}});
                mediaRecorder = new MediaRecorder(stream);
                audioChunks = [];
                mediaRecorder.ondataavailable = e => audioChunks.push(e.data);
                mediaRecorder.onstop = async () => {{
                    const blob = new Blob(audioChunks, {{type:'audio/webm'}});
                    const reader = new FileReader();
                    reader.readAsDataURL(blob);
                    reader.onloadend = async () => {{
                        document.getElementById('status').textContent = '儲存中...';
                        const res = await fetch('/api/sessions/' + SESSION_ID + '/segment', {{
                            method:'POST', headers:{{'Content-Type':'application/json'}},
                            body:JSON.stringify({{audio_data: reader.result.split(',')[1], format:'webm'}})
                        }});
                        const data = await res.json();
                        document.getElementById('status').textContent = data.success ? '✓ 錄音已儲存' : '儲存失敗';
                        await refreshFileList();
                        clearInterval(recordingTimer);
                        stream.getTracks().forEach(t => t.stop());
                    }};
                }};
                mediaRecorder.start(); isRecording = true;
                this.textContent = '⏹️';
                let remaining = 600; // 10分鐘
                recordingTimer = setInterval(() => {{
                    remaining--;
                    const m = Math.floor(remaining/60), s = remaining%60;
                    document.getElementById('status').textContent = '錄音中... 再按停止（還有 ' + m + ':' + s.toString().padStart(2,'0') + '）';
                    if (remaining <= 0) {{ mediaRecorder.stop(); clearInterval(recordingTimer); }}
                }}, 1000);
            }} catch(e) {{ alert('無法存取麥克風：' + e.message); }}
        }} else {{
            mediaRecorder.stop(); isRecording = false; this.textContent = '🎤';
        }}
    }});

    // ========== Upload ==========
    const uploadBox = document.getElementById('uploadBox');
    const fileInput = document.getElementById('fileInput');
    uploadBox.addEventListener('click', () => {{ if (!isProcessing) fileInput.click(); else alert('處理中，無法上傳'); }});
    uploadBox.addEventListener('dragover', e => {{ e.preventDefault(); if(!isProcessing) uploadBox.classList.add('drag-over'); }});
    uploadBox.addEventListener('dragleave', () => uploadBox.classList.remove('drag-over'));
    uploadBox.addEventListener('drop', e => {{
        e.preventDefault(); uploadBox.classList.remove('drag-over');
        if (isProcessing) {{ alert('處理中，無法上傳'); return; }}
        if (e.dataTransfer.files.length > 0) {{ fileInput.files = e.dataTransfer.files; doUpload(fileInput.files); }}
    }});
    fileInput.addEventListener('change', () => {{ if(fileInput.files.length > 0) doUpload(fileInput.files); }});

    async function doUpload(files) {{
        for (const f of files) {{
            if (f.size > 100*1024*1024) {{ alert(f.name + ' 超過 100MB 限制'); return; }}
        }}
        const fd = new FormData();
        for (const f of files) fd.append('files', f);
        const email = document.getElementById('emailInput').value.trim();
        if (email) fd.append('email', email);
        
        document.getElementById('status').textContent = '上傳中...';
        try {{
            const res = await fetch('/api/sessions/' + SESSION_ID + '/upload', {{method:'POST', body:fd}});
            const data = await res.json();
            if (data.success) {{
                document.getElementById('status').textContent = '✓ ' + files.length + ' 個檔案上傳成功';
                await refreshFileList();
            }} else {{
                document.getElementById('status').textContent = '上傳失敗：' + data.error;
            }}
        }} catch(e) {{ document.getElementById('status').textContent = '上傳錯誤'; }}
        fileInput.value = '';
    }}

    // ========== Generate Report ==========
    function startProcessing(mode) {{
        document.getElementById('processingArea').style.display = 'block';
        document.getElementById('genProgressFill').style.width = '2%';
        document.getElementById('genProgressText').textContent = '啟動中...';
        document.getElementById('genResult').innerHTML = '';
        document.getElementById('genSegments').innerHTML = '';
        
        document.getElementById('bgEmailBtn').textContent = '⏳ 處理中...';
        document.getElementById('generateBtn').textContent = '⏳ 處理中...';
        
        genPollTimer = setInterval(pollGenerateStatus, 3000);
        
        if (mode === 'bg') {{
            document.getElementById('genProgressText').textContent = '📧 背景處理中，完成後將自動寄送至信箱（可關閉此頁）';
        }} else {{
            document.getElementById('genProgressText').textContent = '⚡ 即時處理中，請稍候...';
        }}
    }}

    function stopProcessing() {{
        if (genPollTimer) {{ clearInterval(genPollTimer); genPollTimer = null; }}
        isProcessing = false;
        disableInputsDuringProcessing(false);
        document.getElementById('bgEmailBtn').textContent = '📧 背景處理並寄送';
        document.getElementById('generateBtn').textContent = '⚡ 即時處理';
    }}

    // ========== 背景處理並寄送 ==========
    document.getElementById('bgEmailBtn').addEventListener('click', async function() {{
        if (isProcessing) return;
        await refreshFileList();
        const listEl = document.getElementById('audioFileList');
        if (!listEl || listEl.textContent.includes('尚無音檔')) {{
            alert('請先錄音或上傳音檔'); return;
        }}
        
        const email = document.getElementById('emailInput').value.trim();
        if (!email) {{
            alert('請先輸入 Email 再使用背景處理並寄送');
            document.getElementById('emailInput').focus();
            return;
        }}
        
        // Save email to server first by doing a dummy upload with just the email
        // Actually, the upload endpoint also accepts email. Let's set it via a simple API call
        // Use upload with empty files just to set the email... OR we can just pass it in the generate request
        
        // Simpler: set email by making a dummy API call
        // The email is stored on the upload endpoint. Let me just set it.
        // Actually, the /generate endpoint doesn't accept email. Let me store it directly.
        const setRes = await fetch('/api/sessions/' + SESSION_ID + '/set-email', {{
            method:'POST', headers:{{'Content-Type':'application/json'}},
            body:JSON.stringify({{email: email}})
        }});
        const setData = await setRes.json();
        if (!setData.success) {{
            alert('設定 Email 失敗'); return;
        }}
        
        isProcessing = true;
        disableInputsDuringProcessing(true);
        startProcessing('bg');
        
        try {{
            const res = await fetch('/api/sessions/' + SESSION_ID + '/generate', {{
                method:'POST', headers:{{'Content-Type':'application/json'}},
                body:JSON.stringify({{report_type: document.getElementById('reportType').value}})
            }});
            const data = await res.json();
            if (!data.success) {{
                document.getElementById('genResult').innerHTML = '<div class="error">啟動失敗：' + escapeHtml(data.error) + '</div>';
                stopProcessing();
            }}
        }} catch(e) {{
            document.getElementById('genResult').innerHTML = '<div class="error">錯誤：' + escapeHtml(e.message) + '</div>';
            stopProcessing();
        }}
    }});

    // ========== 即時處理 ==========
    document.getElementById('generateBtn').addEventListener('click', async function() {{
        if (isProcessing) return;
        await refreshFileList();
        const listEl = document.getElementById('audioFileList');
        if (!listEl || listEl.textContent.includes('尚無音檔')) {{
            alert('請先錄音或上傳音檔'); return;
        }}
        
        isProcessing = true;
        disableInputsDuringProcessing(true);
        startProcessing('online');
        
        try {{
            const res = await fetch('/api/sessions/' + SESSION_ID + '/generate', {{
                method:'POST', headers:{{'Content-Type':'application/json'}},
                body:JSON.stringify({{report_type: document.getElementById('reportType').value}})
            }});
            const data = await res.json();
            if (!data.success) {{
                document.getElementById('genResult').innerHTML = '<div class="error">啟動失敗：' + escapeHtml(data.error) + '</div>';
                stopProcessing();
            }}
        }} catch(e) {{
            document.getElementById('genResult').innerHTML = '<div class="error">錯誤：' + escapeHtml(e.message) + '</div>';
            stopProcessing();
        }}
    }});

    // ========== Clear ==========
    document.getElementById('clearBtn').addEventListener('click', async () => {{
        if (isProcessing) {{ alert('處理中，無法清空'); return; }}
        const res = await fetch('/api/sessions/' + SESSION_ID + '/clear', {{method:'POST', headers:{{'Content-Type':'application/json'}}}});
        const data = await res.json();
        if (data.success) {{
            document.getElementById('downloadBtn').disabled = true;
            document.getElementById('emailBtn').disabled = true;
            document.getElementById('bgEmailBtn').textContent = '📧 背景處理並寄送';
            document.getElementById('generateBtn').textContent = '⚡ 即時處理';
            document.getElementById('result').innerHTML = '';
            document.getElementById('genResult').innerHTML = '';
            document.getElementById('processingArea').style.display = 'none';
            document.getElementById('audioFileList').innerHTML = '';
            document.getElementById('fileCount').textContent = '(0)';
            document.getElementById('status').textContent = '已清空重置';
        }}
    }});

    // ========== Download & Email ==========
    document.getElementById('downloadBtn').addEventListener('click', () => {{
        window.location.href = '/api/sessions/' + SESSION_ID + '/download';
    }});
    document.getElementById('emailBtn').addEventListener('click', async () => {{
        const email = prompt('請輸入收件者 Email:');
        if (!email) return;
        document.getElementById('emailBtn').textContent = '發送中...'; document.getElementById('emailBtn').disabled = true;
        const res = await fetch('/api/sessions/' + SESSION_ID + '/email', {{
            method:'POST', headers:{{'Content-Type':'application/json'}},
            body:JSON.stringify({{to_email: email}})
        }});
        const data = await res.json();
        alert(data.success ? '✓ 郵件已發送' : '發送失敗：' + data.error);
        document.getElementById('emailBtn').textContent = '📧 手動寄送 Email';
        document.getElementById('emailBtn').disabled = false;
    }});

    // ========== Delete Single Audio File ==========
    async function deleteAudioFile(uid) {{
        if (isProcessing) {{ alert('處理中，無法刪除'); return; }}
        if (!confirm('確定刪除此音檔？')) return;
        try {{
            const res = await fetch('/api/sessions/' + SESSION_ID + '/audio-files/' + uid, {{method:'DELETE'}});
            const data = await res.json();
            if (data.success) {{
                await refreshFileList();
                document.getElementById('status').textContent = '✓ 音檔已刪除';
            }} else {{
                alert('刪除失敗');
            }}
        }} catch(e) {{
            alert('錯誤：' + e.message);
        }}
    }}

    // ========== Drive Folder ==========
    const setDriveFolderBtn = document.getElementById('setDriveFolderBtn');
    const driveFolderInput = document.getElementById('driveFolderInput');
    const driveFolderStatus = document.getElementById('driveFolderStatus');
    
    if (setDriveFolderBtn) {{
        setDriveFolderBtn.addEventListener('click', async () => {{
            const folderId = driveFolderInput.value.trim();
            if (!folderId) {{ driveFolderStatus.innerHTML = '<span class="error">請輸入資料夾 ID</span>'; return; }}
            try {{
                const res = await fetch('/api/sessions/' + SESSION_ID + '/set-drive-folder', {{
                    method: 'POST',
                    headers: {{ 'Content-Type': 'application/json' }},
                    body: JSON.stringify({{ folder_id: folderId }})
                }});
                const data = await res.json();
                if (data.success) {{
                    driveFolderStatus.innerHTML = '<span class="success">✓ Drive 資料夾已設定</span>';
                }} else {{
                    driveFolderStatus.innerHTML = '<span class="error">設定失敗：' + (data.error || '') + '</span>';
                }}
            }} catch(e) {{
                driveFolderStatus.innerHTML = '<span class="error">錯誤：' + e.message + '</span>';
            }}
        }});
    }}

    // ========== Init ==========
    refreshFileList();
    
    // ══════════════════════════════════════════
    // Report Type Management
    // ══════════════════════════════════════════
    const manageTypesBtn = document.getElementById('manageTypesBtn');
    const typeModal = document.getElementById('typeModal');
    const closeModalBtn = document.getElementById('closeModalBtn');
    const typeList = document.getElementById('typeList');
    const newTypeName = document.getElementById('newTypeName');
    const addTypeBtn = document.getElementById('addTypeBtn');
    
    async function loadReportTypes() {{
        try {{
            const [typesRes, tmplRes] = await Promise.all([
                fetch('/api/report-types'),
                fetch('/api/sessions/' + SESSION_ID + '/template')
            ]);
            const typesData = await typesRes.json();
            const tmplData = await tmplRes.json();
            const templates = tmplData.templates || {{}};
            reportType.innerHTML = '';
            typesData.types.forEach(t => {{
                const opt = document.createElement('option');
                opt.value = t.id;
                const hasTmpl = templates[t.id] ? ' 📋' : '';
                opt.textContent = t.name + hasTmpl;
                reportType.appendChild(opt);
            }});
            renderTypeList(typesData.types, templates);
        }} catch(e) {{
            console.error('Failed to load report types:', e);
        }}
    }}
    
    function renderTypeList(types, templates) {{
        templates = templates || {{}};
        typeList.innerHTML = types.map(t => {{
            const isDefault = ['general','medical','meeting','swallow','ent'].includes(t.id);
            const hasTmpl = templates[t.id] ? ' 📋' : '';
            return '<div class="type-list-item">' +
                '<span class="type-name">' + escapeHtml(t.name) + hasTmpl + '</span>' +
                (isDefault ? '<span class="type-badge">預設</span>' : '<button class="type-delete-btn" data-id="' + t.id + '" title="刪除">✕</button>') +
            '</div>';
        }}).join('');
        document.querySelectorAll('.type-delete-btn').forEach(btn => {{
            btn.addEventListener('click', async () => {{
                const typeId = btn.dataset.id;
                const name = btn.parentElement.querySelector('.type-name').textContent;
                if (!confirm('確定刪除「' + name + '」？')) return;
                try {{
                    const res = await fetch('/api/report-types/' + typeId, {{ method: 'DELETE' }});
                    const data = await res.json();
                    if (data.success) await loadReportTypes();
                    else alert('刪除失敗：' + (data.detail || data.error || '未知錯誤'));
                }} catch(e) {{ alert('刪除失敗：' + e.message); }}
            }});
        }});
    }}
    
    if (manageTypesBtn) {{
        manageTypesBtn.addEventListener('click', () => {{
            typeModal.style.display = 'flex';
            loadReportTypes();
        }});
    }}
    if (closeModalBtn) {{
        closeModalBtn.addEventListener('click', () => {{ typeModal.style.display = 'none'; }});
    }}
    if (typeModal) {{
        typeModal.addEventListener('click', (e) => {{
            if (e.target === typeModal) typeModal.style.display = 'none';
        }});
    }}
    if (addTypeBtn) {{
        addTypeBtn.addEventListener('click', async () => {{
            const name = newTypeName.value.trim();
            if (!name) {{ alert('請輸入報告類型名稱'); return; }}
            try {{
                const res = await fetch('/api/report-types', {{
                    method: 'POST',
                    headers: {{ 'Content-Type': 'application/json' }},
                    body: JSON.stringify({{ name: name }})
                }});
                const data = await res.json();
                if (data.success) {{ newTypeName.value = ''; await loadReportTypes(); }}
                else alert('新增失敗：' + (data.detail || data.error || '未知錯誤'));
            }} catch(e) {{ alert('新增失敗：' + e.message); }}
        }});
        newTypeName.addEventListener('keydown', (e) => {{
            if (e.key === 'Enter') addTypeBtn.click();
        }});
    }}
    loadReportTypes();
    </script>
</body></html>"""
    return HTMLResponse(content=html)


@app.get("/new")
async def new_session(request: Request):
    if not validate_session(request):
        return RedirectResponse(url="/login", status_code=302)
    new_id = get_or_create_session()
    response = RedirectResponse(url="/dashboard", status_code=302)
    response.set_cookie(key="session_id", value=new_id, httponly=True, samesite="lax")
    return response


# =============================================================================
# Other API Endpoints
# =============================================================================

@app.post("/api/sessions/{session_id}/set-email")
async def set_email(session_id: str, request: Request):
    if not validate_session(request) or get_session_id(request) != session_id:
        raise HTTPException(status_code=401, detail="Not authenticated")
    if session_id not in sessions:
        raise HTTPException(status_code=404, detail="Session not found")
    data = await request.json()
    email = data.get("email", "")
    if not email:
        return {"success": False, "error": "Email 不得為空"}
    sessions[session_id]["auto_email"] = email
    sessions[session_id]["auto_email_sent"] = False
    sessions[session_id]["auto_email_message"] = None
    sessions[session_id]["auto_email_error"] = None
    return {"success": True, "email": email}


@app.post("/api/sessions/{session_id}/set-drive-folder")
async def set_drive_folder(session_id: str, request: Request):
    if not validate_session(request) or get_session_id(request) != session_id:
        raise HTTPException(status_code=401, detail="Not authenticated")
    if session_id not in sessions:
        raise HTTPException(status_code=404, detail="Session not found")
    data = await request.json()
    folder_id = data.get("folder_id", "").strip()
    sessions[session_id]["drive_folder_id"] = folder_id
    sessions[session_id]["drive_upload_result"] = None
    return {"success": True, "folder_id": folder_id}


@app.get("/api/sessions/{session_id}/template")
async def get_template_status(session_id: str, request: Request):
    if not validate_session(request) or get_session_id(request) != session_id:
        raise HTTPException(status_code=401, detail="Not authenticated")
    # Check which report types have templates
    templates = {}
    for f in TEMPLATES_DIR.glob("*.docx"):
        type_id = f.stem  # filename without .docx
        if type_id == "template":
            continue  # skip generic fallback
        templates[type_id] = True
    return {"exists": (TEMPLATES_DIR / "template.docx").exists(), "templates": templates}

@app.post("/api/sessions/{session_id}/template")
async def upload_template(session_id: str, request: Request, file: UploadFile = File(...), report_type: str = Form("template")):
    if not validate_session(request) or get_session_id(request) != session_id:
        raise HTTPException(status_code=401, detail="Not authenticated")
    if not file.filename.endswith('.docx'):
        raise HTTPException(status_code=400, detail="Only .docx files are accepted")
    content = await file.read()
    
    # Save locally per report type
    local_path = TEMPLATES_DIR / f"{report_type}.docx"
    with open(local_path, "wb") as f:
        f.write(content)
    
    # Also save as generic fallback
    with open(TEMPLATES_DIR / "template.docx", "wb") as f:
        f.write(content)
    
    # Upload to Google Drive templates folder
    if MATON_API_KEY:
        drive_filename = f"{report_type}_template.docx"
        success, result = _upload_binary_to_drive(content, drive_filename, DRIVE_TEMPLATES_FOLDER_ID)
        drive_status = "✓ 已同步到雲端" if success else f"⚠️ 雲端同步失敗：{result}"
    else:
        drive_status = "⚠️ 未設定 Maton API，僅儲存本地"
    
    return {"success": True, "filename": file.filename, "report_type": report_type, "drive_status": drive_status}

@app.post("/api/auth/login")
async def login(request: Request):
    data = await request.json()
    if data.get("password") == ACCESS_PASSWORD:
        session_id = get_or_create_session()
        response = JSONResponse({"success": True})
        response.set_cookie(key="session_id", value=session_id, httponly=True, samesite="lax")
        return response
    raise HTTPException(status_code=401, detail="Invalid password")

@app.post("/api/sessions/{session_id}/clear")
async def clear_session(session_id: str, request: Request):
    if not validate_session(request) or get_session_id(request) != session_id:
        raise HTTPException(status_code=401, detail="Not authenticated")
    if session_id not in sessions:
        raise HTTPException(status_code=404, detail="Session not found")
    # Reset session state while keeping session alive
    s = sessions[session_id]
    s["segments"] = []
    s["audio_files"] = {}
    s["generated_docx"] = None
    s["auto_email"] = ""
    s["auto_email_sent"] = False
    s["auto_email_message"] = None
    s["auto_email_error"] = None
    s["processing"] = False
    s["processing_done"] = False
    s["processing_error"] = None
    s["processing_progress"] = None
    return {"success": True}

@app.get("/api/sessions/{session_id}/download")
async def download_report(session_id: str, request: Request):
    if not validate_session(request) or get_session_id(request) != session_id:
        raise HTTPException(status_code=401, detail="Not authenticated")
    if session_id not in sessions:
        raise HTTPException(status_code=404, detail="Session not found")
    session = sessions[session_id]
    if session.get("generated_docx"):
        today = datetime.now().strftime('%Y%m%d')
        return Response(content=session["generated_docx"],
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=MCH_Report_{today}.docx"})
    elif session.get("segments"):
        return Response(content="\n\n".join([s["transcription"] for s in session["segments"]]),
            media_type="text/plain; charset=utf-8")
    raise HTTPException(status_code=404, detail="No report to download")

@app.delete("/api/sessions/{session_id}/audio-files/{uid}")
async def delete_audio_file(session_id: str, uid: str, request: Request):
    if not validate_session(request) or get_session_id(request) != session_id:
        raise HTTPException(status_code=401, detail="Not authenticated")
    if session_id not in sessions:
        raise HTTPException(status_code=404, detail="Session not found")
    
    audio_files = sessions[session_id].get("audio_files", {})
    if uid not in audio_files:
        raise HTTPException(status_code=404, detail="Audio file not found")
    
    # Delete the actual file on disk
    filepath = audio_files[uid].get("filepath")
    if filepath and Path(filepath).exists():
        try:
            Path(filepath).unlink()
        except Exception as e:
            logging.warning(f"Failed to delete file {filepath}: {e}")
    
    # Remove from session
    del audio_files[uid]
    
    return {"success": True}


@app.post("/api/sessions/{session_id}/email")
async def email_report(session_id: str, request: Request, req: EmailRequest):
    if not validate_session(request) or get_session_id(request) != session_id:
        raise HTTPException(status_code=401, detail="Not authenticated")
    if session_id not in sessions:
        raise HTTPException(status_code=404, detail="Session not found")
    session = sessions[session_id]
    if not session.get("generated_docx") and not session.get("segments"):
        return {"success": False, "error": "無報告可發送"}
    success, msg = _send_email_sync(session, req.to_email)
    return {"success": success, "message" if success else "error": msg}

@app.get("/health")
async def health():
    return {"status": "healthy", "service": "MCH Assistant", "version": "2.1.0"}

@app.get("/api/status")
async def api_status():
    return {"service": "MCH Assistant", "version": "2.1.0", "status": "running", "sessions": len(sessions)}

# =============================================================================
# Dysphagia Screening Report Upload
# =============================================================================

def _upload_to_drive(content: str, filename: str) -> tuple:
    """Upload text report to Google Drive via Maton API. Returns (success, fileId or error)."""
    import urllib.request, json

    file_data = content.encode('utf-8')
    file_size = len(file_data)

    # Step 1: Start resumable upload session
    meta = json.dumps({
        "name": filename,
        "parents": [DRIVE_FOLDER_ID]
    }).encode()

    req1 = urllib.request.Request(
        DRIVE_UPLOAD_URL + "?uploadType=resumable",
        data=meta,
        headers={
            "Authorization": f"Bearer {MATON_API_KEY}",
            "Content-Type": "application/json; charset=UTF-8",
            "X-Upload-Content-Type": "text/plain; charset=utf-8",
            "X-Upload-Content-Length": str(file_size)
        },
        method="POST"
    )

    try:
        resp1 = urllib.request.urlopen(req1, timeout=15)
        location = resp1.getheader("Location")
    except Exception as e:
        return False, f"無法建立上傳連線：{str(e)}"

    # Step 2: PUT file content to session URL
    req2 = urllib.request.Request(
        location,
        data=file_data,
        headers={
            "Content-Type": "text/plain; charset=utf-8",
            "Content-Length": str(file_size)
        },
        method="PUT"
    )

    try:
        resp2 = urllib.request.urlopen(req2, timeout=30)
        result = json.loads(resp2.read())
        file_id = result.get("id", "unknown")
    except Exception as e:
        return False, f"上傳失敗：{str(e)}"

    # Step 3: Verify by downloading
    try:
        verify_req = urllib.request.Request(
            f"https://gateway.maton.ai/google-drive/drive/v3/files/{file_id}?alt=media",
            headers={"Authorization": f"Bearer {MATON_API_KEY}"}
        )
        with urllib.request.urlopen(verify_req, timeout=20) as vresp:
            verified_size = len(vresp.read())
            if verified_size != file_size:
                return False, f"檔案驗證失敗（大小不符）"
    except Exception as e:
        return False, f"檔案驗證失敗：{str(e)}"

    return True, file_id


def _upload_binary_to_drive(file_bytes: bytes, filename: str, folder_id: str = None) -> tuple:
    """Upload binary file (e.g. .docx) to Google Drive via Maton API using multipart upload."""
    import urllib.request
    
    target_folder = folder_id or DRIVE_FOLDER_ID
    boundary = "----MCHBoundary" + uuid.uuid4().hex[:16]
    
    # Metadata JSON part
    meta_part = json.dumps({
        "name": filename,
        "parents": [target_folder]
    })
    
    # Build multipart body
    body_parts = []
    body_parts.append(f"--{boundary}\r\n".encode())
    body_parts.append(b"Content-Type: application/json; charset=UTF-8\r\n\r\n")
    body_parts.append(meta_part.encode("utf-8"))
    body_parts.append(f"\r\n--{boundary}\r\n".encode())
    body_parts.append(b"Content-Type: application/vnd.openxmlformats-officedocument.wordprocessingml.document\r\n")
    body_parts.append(b"Content-Transfer-Encoding: binary\r\n\r\n")
    body_parts.append(file_bytes)
    body_parts.append(f"\r\n--{boundary}--\r\n".encode())
    
    body = b"".join(body_parts)
    
    req = urllib.request.Request(
        DRIVE_UPLOAD_URL + "?uploadType=multipart",
        data=body,
        headers={
            "Authorization": f"Bearer {MATON_API_KEY}",
            "Content-Type": f"multipart/related; boundary={boundary}",
            "Content-Length": str(len(body))
        },
        method="POST"
    )
    
    try:
        resp = urllib.request.urlopen(req, timeout=30)
        result = json.loads(resp.read())
        file_id = result.get("id", "unknown")
        return True, file_id
    except Exception as e:
        return False, f"模板上傳失敗：{str(e)}"


DRIVE_FILES_API = "https://gateway.maton.ai/google-drive/drive/v3/files"


def _list_drive_templates(folder_id: str) -> list:
    """List all template files (*_template.docx) in the Drive folder."""
    import urllib.request
    query = f"'{folder_id}' in parents and name contains '_template' and trashed = false"
    url = f"{DRIVE_FILES_API}?q={urllib.parse.quote(query)}&fields=files(id,name,mimeType,size)"
    req = urllib.request.Request(url, headers={"Authorization": f"Bearer {MATON_API_KEY}"})
    try:
        resp = urllib.request.urlopen(req, timeout=15)
        data = json.loads(resp.read())
        return data.get("files", [])
    except Exception as e:
        logging.warning(f"Drive list error: {e}")
        return []


def _download_drive_file(file_id: str) -> Optional[bytes]:
    """Download a file from Google Drive by ID."""
    import urllib.request
    url = f"{DRIVE_FILES_API}/{file_id}?alt=media"
    req = urllib.request.Request(url, headers={"Authorization": f"Bearer {MATON_API_KEY}"})
    try:
        resp = urllib.request.urlopen(req, timeout=30)
        return resp.read()
    except Exception as e:
        logging.warning(f"Drive download error (file {file_id}): {e}")
        return None


def _sync_templates_from_drive():
    """Sync templates from Google Drive to local templates directory."""
    if not MATON_API_KEY:
        logging.warning("Maton API Key not configured, skipping Drive template sync")
        return
    
    logging.info("Syncing templates from Google Drive...")
    files = _list_drive_templates(DRIVE_TEMPLATES_FOLDER_ID)
    if not files:
        logging.info("No templates found in Drive folder")
        return
    
    synced = 0
    for f in files:
        name = f.get("name", "")
        file_id = f.get("id", "")
        if not name.endswith("_template.docx") or not file_id:
            continue
        # Extract report type: "swallow_template.docx" → "swallow"
        type_id = name.replace("_template.docx", "")
        if not type_id:
            continue
        
        data = _download_drive_file(file_id)
        if data is None:
            continue
        
        # Save locally
        local_path = TEMPLATES_DIR / f"{type_id}.docx"
        local_path.write_bytes(data)
        logging.info(f"  ✓ Synced template: {name} → templates/{type_id}.docx")
        synced += 1
    
    logging.info(f"Template sync complete: {synced} files synced")


@app.post("/api/dysphagia/upload")
async def dysphagia_upload(req: DysphagiaUploadRequest, request: Request):
    """Receive dysphagia screening report and upload to Google Drive."""
    # Check if MATON_API_KEY is configured
    if not MATON_API_KEY:
        return JSONResponse(
            status_code=500,
            content={"success": False, "error": "Maton API Key 未設定，請聯繫系統管理員"}
        )

    # Build filename: 吞嚥篩檢報告_[姓名]_[日期].txt
    safe_name = re.sub(r'[\\/:*?"<>|]', '_', req.patientName or '未命名')
    filename = f"吞嚥篩檢報告_{safe_name}_{req.date}.txt"

    # Upload to Google Drive
    success, result = _upload_to_drive(req.report, filename)

    if success:
        return {"success": True, "fileId": result, "filename": filename}
    else:
        return JSONResponse(status_code=500, content={"success": False, "error": result})


@app.post("/api/dysphagia/email")
async def dysphagia_email(req: DysphagiaUploadRequest, request: Request):
    """Send dysphagia screening report via email using Maton Gmail API."""
    if not MATON_API_KEY:
        return JSONResponse(
            status_code=500,
            content={"success": False, "error": "Maton API Key 未設定，請聯繫系統管理員"}
        )

    if not req.to:
        return JSONResponse(
            status_code=400,
            content={"success": False, "error": "請提供收件者 Email"}
        )

    import urllib.request

    # Build email metadata
    safe_name = re.sub(r'[\\/:*?"<>|]', '_', req.patientName or '未命名')
    subject = f"【門諾醫院】吞嚥篩檢報告 - {safe_name} ({req.date})"

    # Build HTML body
    html_body = f"""
    <html>
    <body style="font-family: 'Noto Sans TC', Arial, sans-serif; max-width: 700px; margin: 0 auto; padding: 20px;">
      <div style="background: linear-gradient(135deg, #2c6e9e, #1a4a6b); color: white; padding: 24px; border-radius: 12px 12px 0 0; text-align: center;">
        <h2 style="margin: 0;">🏥 門諾醫院 吞嚥篩檢報告</h2>
        <p style="margin: 8px 0 0; opacity: 0.9;">Mennonite Christian Hospital - Dysphagia Screening Report</p>
      </div>
      <div style="background: #f8fbfd; padding: 24px; border: 1px solid #d0d9e0; border-top: none; border-radius: 0 0 12px 12px;">
        <p style="color: #555; font-size: 14px;">📋 <strong>受檢者：</strong>{safe_name} &nbsp;&nbsp; 📅 <strong>篩檢日期：</strong>{req.date}</p>
        <hr style="border: none; border-top: 1px solid #d0d9e0; margin: 16px 0;">
        <pre style="white-space: pre-wrap; word-wrap: break-word; font-size: 14px; line-height: 1.7; color: #333; font-family: inherit;">{req.report}</pre>
        <hr style="border: none; border-top: 1px solid #d0d9e0; margin: 16px 0;">
        <p style="font-size: 12px; color: #888; text-align: center;">
          本報告由門諾醫院吞嚥團隊 AI 助理自動產生，僅供參考，不作為診斷依據。<br>
          如有吞嚥問題請諮詢專業醫療人員。<br>
          <em>This report is automatically generated and for reference only. Please consult a healthcare professional for any medical concerns.</em>
        </p>
      </div>
    </body>
    </html>
    """

    # Build plain text body
    plain_body = f"""門諾醫院 吞嚥篩檢報告
====================
受檢者：{safe_name}
篩檢日期：{req.date}

{req.report}

---
本報告由門諾醫院吞嚥團隊 AI 助理自動產生，僅供參考，不作為診斷依據。
如有吞嚥問題請諮詢專業醫療人員。
"""

    # Construct RFC 2822 email
    from email.mime.multipart import MIMEMultipart
    from email.mime.text import MIMEText
    import email.utils

    msg = MIMEMultipart('alternative')
    msg['Subject'] = subject
    msg['From'] = "mchswallow@gmail.com"
    msg['To'] = req.to
    msg['Date'] = email.utils.formatdate(localtime=True)

    # Encode properly for SMTP
    import base64
    html_encoded = base64.b64encode(html_body.encode('utf-8')).decode('utf-8')
    plain_encoded = base64.b64encode(plain_body.encode('utf-8')).decode('utf-8')

    # Use Maton Gmail API to send
    gmail_send_url = "https://gateway.maton.ai/google-mail/gmail/v1/users/me/messages/send"

    # Construct the raw email (RFC 5322 format with base64 encoded parts)
    import email.generator
    import io as _io

    msg.attach(MIMEText(plain_body, 'plain', 'utf-8'))
    msg.attach(MIMEText(html_body, 'html', 'utf-8'))

    # Get raw email bytes
    buf = _io.BytesIO()
    gen = email.generator.BytesGenerator(buf)
    gen.flatten(msg)
    raw_email_b64 = base64.urlsafe_b64encode(buf.getvalue()).decode('utf-8')

    payload = json.dumps({"raw": raw_email_b64})

    req_gmail = urllib.request.Request(
        gmail_send_url,
        data=payload.encode('utf-8'),
        headers={
            "Authorization": f"Bearer {MATON_API_KEY}",
            "Content-Type": "application/json"
        },
        method="POST"
    )

    try:
        resp = urllib.request.urlopen(req_gmail, timeout=30)
        result_data = json.loads(resp.read())
        thread_id = result_data.get('threadId', '')
        return {"success": True, "messageId": result_data.get('id', ''), "threadId": thread_id}
    except Exception as e:
        return JSONResponse(status_code=500, content={"success": False, "error": f"寄送失敗：{str(e)}"})

# =============================================================================
# Report Types CRUD API
# =============================================================================

@app.get("/api/report-types")
async def api_list_report_types():
    """列出所有報告類型"""
    types = [{"id": k, "name": v} for k, v in sorted(report_types_store.items())]
    return {"success": True, "types": types}

@app.post("/api/report-types")
async def api_add_report_type(req: Request):
    """新增報告類型（只需提供 name，自動生成 id）"""
    body = await req.json()
    name = body.get("name", "").strip()
    if not name:
        raise HTTPException(status_code=400, detail="name 為必填")
    
    # Generate unique id from name (slug) + random suffix if needed
    type_id = re.sub(r'[^a-zA-Z0-9\u4e00-\u9fff]', '_', name).lower()[:20]
    if type_id in report_types_store:
        # Append random suffix for uniqueness
        type_id = f"{type_id}_{uuid.uuid4().hex[:4]}"
    
    report_types_store[type_id] = name
    save_report_types(report_types_store)
    return {"success": True, "id": type_id, "name": name}

@app.put("/api/report-types/{type_id}")
async def api_rename_report_type(type_id: str, req: Request):
    """重新命名報告類型"""
    if type_id not in report_types_store:
        raise HTTPException(status_code=404, detail="報告類型不存在")
    body = await req.json()
    name = body.get("name", "").strip()
    if not name:
        raise HTTPException(status_code=400, detail="name 為必填")
    report_types_store[type_id] = name
    save_report_types(report_types_store)
    return {"success": True, "id": type_id, "name": name}

@app.delete("/api/report-types/{type_id}")
async def api_delete_report_type(type_id: str):
    """刪除報告類型"""
    if type_id not in report_types_store:
        raise HTTPException(status_code=404, detail="報告類型不存在")
    if type_id in DEFAULT_REPORT_TYPES:
        raise HTTPException(status_code=400, detail="無法刪除預設報告類型")
    del report_types_store[type_id]
    save_report_types(report_types_store)
    return {"success": True}


@app.post("/api/report-types/sync-templates")
async def api_sync_templates():
    """手動觸發從 Google Drive 同步模板"""
    try:
        _sync_templates_from_drive()
        synced = []
        for f in TEMPLATES_DIR.glob("*.docx"):
            if f.stem != "template":
                synced.append(f.stem)
        return {"success": True, "synced": synced, "count": len(synced)}
    except Exception as e:
        return JSONResponse(status_code=500, content={"success": False, "error": str(e)})


# =============================================================================
# Run
# =============================================================================

if __name__ == "__main__":
    import uvicorn
    port = int(os.getenv("PORT", "8080"))
    uvicorn.run(app, host="0.0.0.0", port=port)
